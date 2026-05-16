"""
Add generated figures to the thesis document at the appropriate positions.
"""

from docx import Document
from docx.shared import Inches
import os

doc = Document(r'C:\Users\HP\Downloads\RAG_Medical_QA_Thesis_LJMU_Final_Updated.docx')

figures = {
    'Figure 5.1:': 'reports/figures/figure_5_1_faithfulness.png',
    'Figure 5.2:': 'reports/figures/figure_5_2_precision_recall.png',
    'Figure 5.3:': 'reports/figures/figure_5_3_hallucination.png',
    'Figure 5.4:': 'reports/figures/figure_5_4_groundedness_correctness.png',
    'Figure 5.5:': 'reports/figures/figure_5_5_per_type.png',
    'Figure 5.6:': 'reports/figures/figure_5_6_safety.png',
    'Figure 5.7:': 'reports/figures/figure_5_7_failures.png',
    'Figure 5.8:': 'reports/figures/figure_5_8_latency.png',
    'Figure 5.9:': 'reports/figures/figure_5_9_correlation.png',
}

# Process in reverse order to maintain indices
for i, para in enumerate(doc.paragraphs):
    para_text = para.text.strip()
    for fig_prefix, fig_path in figures.items():
        if para_text.startswith(fig_prefix):
            if os.path.exists(fig_path):
                # Add image before the caption paragraph
                run = para.insert_paragraph_before().add_run()
                run.add_picture(fig_path, width=Inches(5.5))
                para.insert_paragraph_before().alignment = 1  # Center align
                print(f"Added {fig_path} before '{para_text[:40]}...'")
            else:
                print(f"WARNING: {fig_path} not found")
            break

output_path = r'C:\Users\HP\Downloads\RAG_Medical_QA_Thesis_LJMU_Final_With_Figures.docx'
doc.save(output_path)
print(f"\nThesis with figures saved to: {output_path}")
