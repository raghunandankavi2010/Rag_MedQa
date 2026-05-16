"""
Add page numbers to thesis document.
In Word, users should right-click TOC and select 'Update Field' to refresh page numbers.
This script sets up proper headers/footers with page number fields.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document(r'C:\Users\HP\Downloads\RAG_Medical_QA_Thesis_LJMU_FINAL.docx')

# Helper to add page number field
def add_page_number(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(80, 80, 80)

def add_num_pages(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "NUMPAGES"
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(80, 80, 80)

# Add page numbers to all sections
for section in doc.sections:
    footer = section.footer
    footer.is_linked_to_previous = False
    
    # Clear existing footer content
    for paragraph in footer.paragraphs:
        paragraph.clear()
    
    # Add centered page number
    if len(footer.paragraphs) == 0:
        footer.add_paragraph()
    
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add "Page X of Y" format
    run1 = paragraph.add_run("Page ")
    run1.font.size = Pt(10)
    run1.font.color.rgb = RGBColor(80, 80, 80)
    
    add_page_number(paragraph)
    
    run2 = paragraph.add_run(" of ")
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(80, 80, 80)
    
    add_num_pages(paragraph)

# Add a note at the beginning about updating TOC
# Find the TOC paragraph
toc_para = None
for i, para in enumerate(doc.paragraphs):
    if "Right-click the table of contents" in para.text:
        toc_para = para
        break

if toc_para:
    toc_para.text = "Right-click the table of contents and select 'Update Field' → 'Update entire table' to refresh all page numbers after opening in Microsoft Word."

output_path = r'C:\Users\HP\Downloads\RAG_Medical_QA_Thesis_LJMU_COMPLETE.docx'
doc.save(output_path)
print(f"Thesis with page numbers saved to: {output_path}")
print("IMPORTANT: Open in Microsoft Word and right-click the TOC to update page numbers.")
