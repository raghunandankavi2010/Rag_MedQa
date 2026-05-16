"""
Add generated figures to the thesis document ONLY in Chapter 5 text.
"""

from docx import Document
from docx.shared import Inches
import os

doc = Document(r'C:\Users\HP\Downloads\RAG_Medical_QA_Thesis_LJMU_Final_Updated.docx')

figures = {
    'Figure 5.1: RAGAS Faithfulness Comparison': 'reports/figures/figure_5_1_faithfulness.png',
    'Figure 5.2: RAGAS Context Precision and Recall': 'reports/figures/figure_5_2_precision_recall.png',
    'Figure 5.3: DeepEval Hallucination Rate': 'reports/figures/figure_5_3_hallucination.png',
    'Figure 5.4: DeepEval Groundedness and Correctness': 'reports/figures/figure_5_4_groundedness_correctness.png',
    'Figure 5.5: Per-Question-Type Faithfulness': 'reports/figures/figure_5_5_per_type.png',
    'Figure 5.6: Safety Compliance Scores': 'reports/figures/figure_5_6_safety.png',
    'Figure 5.7: Failure Pattern Distribution': 'reports/figures/figure_5_7_failures.png',
    'Figure 5.8: Latency Comparison': 'reports/figures/figure_5_8_latency.png',
    'Figure 5.9: Correlation Between Faithfulness and Hallucination': 'reports/figures/figure_5_9_correlation.png',
}

# Track which figures we've already added to avoid duplicates
added = set()

for i, para in enumerate(doc.paragraphs):
    para_text = para.text.strip()
    for fig_caption, fig_path in figures.items():
        if fig_caption in para_text and fig_caption not in added:
            if os.path.exists(fig_path):
                # Add image before the caption paragraph
                new_para = para.insert_paragraph_before()
                run = new_para.add_run()
                run.add_picture(fig_path, width=Inches(5.5))
                new_para.alignment = 1  # Center align
                added.add(fig_caption)
                print(f"Added {fig_path}")
            else:
                print(f"WARNING: {fig_path} not found")
            break

output_path = r'C:\Users\HP\Downloads\RAG_Medical_QA_Thesis_LJMU_FINAL.docx'
doc.save(output_path)
print(f"\nFinal thesis saved to: {output_path}")
