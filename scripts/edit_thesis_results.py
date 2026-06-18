#!/usr/bin/env python3
"""Replace fabricated Chapter 5 / Abstract content with the REAL results.

Reads reports/thesis_stats.json (+ raw CSV for median latency) and rewrites the
result tables, prose, captions, methodology config and ANOVA in the thesis
.docx so that every reported number is the genuine output of the pilot run
(n = 50 per pipeline, N = 250). Passive academic voice throughout.
"""
import json
import docx
import pandas as pd

PATH = "/Users/raghunandan.k/Downloads/RaghunandanKavi_PN1196933_RAG_Medical_QA_FINAL.docx"
S = json.load(open("reports/thesis_stats.json"))
PP = S["per_pipeline"]
ORDER = ["P1_Vanilla_LLM", "P2_Standard_RAG", "P3_Multi_Query_Expansion",
         "P4_Hybrid_Retrieval", "P5_Query_Reformulation"]
raw = pd.read_csv("reports/pipeline_comparison_raw.csv")
MEDLAT = {p: raw[raw.pipeline == p]["latency_ms"].median() / 1000.0 for p in ORDER}


def m(pid, metric, key="mean"):
    return PP[pid][metric][key]


def f3(x):
    return "N/A" if x is None else f"{x:.3f}"


# ── helpers ──────────────────────────────────────────────────────────────────
def set_text(p, text):
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


def set_cell(cell, text):
    p = cell.paragraphs[0]
    set_text(p, text)
    for extra in cell.paragraphs[1:]:
        set_text(extra, "")


def replace_prefix(paras, prefix, newtext):
    for p in paras:
        if p.text.strip().startswith(prefix):
            set_text(p, newtext)
            return True
    print("  [WARN] prefix not found:", prefix[:50])
    return False


d = docx.Document(PATH)
paras = d.paragraphs
T = d.tables

# ── GLOBAL config / sample-size string fixes (scattered mentions) ────────────
GLOBAL = [
    ("text-embedding-3-small", "text-embedding-3-large"),
    ("1,536", "3,072"), ("1536", "3072"),
    ("fixed chunk size of 400 tokens with 50-token overlap", "fixed chunk size of 1,200 characters with 150-character overlap"),
    ("400 tokens per chunk with 50-token overlap", "1,200 characters per chunk with 150-character overlap"),
    ("400 tokens with 50-token overlap", "1,200 characters with 150-character overlap"),
    ("chunk size of 400 tokens", "chunk size of 1,200 characters"),
    ("400-token chunks", "1,200-character chunks"),
    ("400-token size", "1,200-character size"),
    ("50-token overlap", "150-character overlap"),
    ("400 tokens", "1,200 characters"),
    ("sample of 100 questions", "sample of 50 questions"),
    ("100 questions per pipeline", "50 questions per pipeline"),
    ("sample of 100", "sample of 50"),
    ("of 100 questions", "of 50 questions"),
    ("100 consumer-health questions", "50 consumer-health questions"),
    ("five groups of 100 observations each (N = 500)", "five groups of 50 observations each (N = 250)"),
    ("n = 100 per group", "n = 50 per group"),
    ("(n = 100)", "(n = 50)"),
    ("across the pilot evaluation sample of 100 questions", "across the pilot evaluation sample of 50 questions"),
]
for p in paras:
    if not p.text.strip():
        continue
    for a, b in GLOBAL:
        if a in p.text:
            for r in p.runs:
                if a in r.text:
                    r.text = r.text.replace(a, b)
            # fall back if split across runs
            if a in p.text and len(p.runs) > 1:
                set_text(p, p.text.replace(a, b))

# ── FULL-PARAGRAPH rewrites (real numbers, passive voice) ────────────────────
R = {}

R["The pipelines were evaluated in terms of grounding quality"] = (
    "The pipelines were evaluated in terms of grounding quality, assessed using the RAGAS metrics of "
    "faithfulness, answer relevance, context precision, and context recall, and in terms of output "
    "reliability, assessed using the DeepEval hallucination and answer-relevance metrics together with "
    "ROUGE overlap against the reference answers and a heuristic safety-orientation score. The MedQuAD "
    "dataset, which contains 47,457 question-answer pairs derived from twelve National Institutes of Health "
    "sources, provided the basis for the retrieval corpus and the evaluation benchmark; a cleaned subset of "
    "14,828 pairs (26,083 indexed chunks) formed the retrieval corpus for this study. The evaluation "
    "reported here was conducted at a pilot scale, on a stratified sample of 50 questions per pipeline; a "
    "full-scale evaluation of 1,000 questions is identified as a priority for future work.")

R["Across the pilot evaluation, hybrid retrieval achieved"] = (
    "Across the pilot evaluation, all four retrieval-augmented pipelines achieved higher faithfulness "
    "(0.485 to 0.614) than the non-retrieval baseline (0.381) and lower hallucination rates (0.248 to "
    "0.283, against 0.320 for the baseline), indicating that grounding generation in retrieved evidence "
    "reduces unsupported content. The highest faithfulness was recorded by multi-query expansion (0.614) "
    "and the highest context precision by standard RAG (0.725), whereas hybrid retrieval did not outperform "
    "the other retrieval strategies on the primary grounding metrics. A one-way analysis of variance on "
    "faithfulness was not statistically significant at this sample size (F(4, 245) = 2.08, p = 0.084), so "
    "the differences are reported as indicative rather than confirmatory. A trade-off was observed between "
    "grounding and fluency: the baseline produced the most relevant and lexically complete answers because "
    "it always responds, whereas the retrieval pipelines were more conservative and occasionally declined to "
    "answer when the retrieved evidence was insufficient. These results are indicative of pilot-scale "
    "performance and require confirmation at full scale.")

R["The OpenAI GPT-4o-mini model, accessed through an API, was selected for generation"] = (
    "The OpenAI GPT-4o-mini model, accessed through an API, was selected for generation, and the "
    "text-embedding-3-large model, also accessed through an API, was selected for embedding. This "
    "combination provides a balance between rapid, reliable experimentation and a strong semantic "
    "representation of medical text. A fixed chunk size of 1,200 characters with 150-character overlap was "
    "used across all five retrieval pipelines.")

R["The vanilla LLM baseline achieves a mean faithfulness of 0.42"] = (
    "The vanilla baseline records the lowest mean faithfulness, 0.381, which is consistent with its reliance "
    "on parametric knowledge alone and with the absence of any retrieved context against which its claims "
    "can be grounded; the large standard deviation (0.472) reflects substantial variation across questions. "
    "All four retrieval-augmented pipelines record higher mean faithfulness: standard RAG 0.494, hybrid "
    "retrieval 0.485, query reformulation 0.512, and multi-query expansion 0.614, the highest observed. The "
    "ordering indicates that broadening query coverage through multi-query expansion retrieved evidence in "
    "which the generator could more consistently ground its claims, whereas hybrid retrieval conferred no "
    "advantage over the simpler dense-retrieval configuration on this sample. As shown in Table 5.1, the "
    "95 per cent confidence intervals overlap substantially, so the between-pipeline differences are modest "
    "relative to the within-pipeline variance; this is examined formally in Section 5.8.")

R["The standard RAG pipeline achieves 0.74 precision and 0.71 recall"] = (
    "Among the retrieval pipelines, the highest context precision is recorded by standard RAG (0.725), "
    "followed by multi-query expansion (0.719), query reformulation (0.656), and hybrid retrieval (0.642). "
    "The highest context recall is recorded by multi-query expansion (0.578), followed by hybrid retrieval "
    "(0.532), standard RAG (0.512), and query reformulation (0.468). Because the vanilla baseline performs "
    "no retrieval, context precision and recall are not applicable to it and are reported as such in "
    "Table 5.2. Hypothesis 3, which predicted that multi-query expansion would achieve the highest recall "
    "and the lowest precision among the retrieval pipelines, is only partially supported: multi-query "
    "expansion does achieve the highest recall, consistent with its broader query coverage, but it does not "
    "record the lowest precision, which is observed instead for hybrid retrieval. The pattern suggests that "
    "query expansion retrieved additional relevant passages without a marked loss of precision on this "
    "sample, whereas the reciprocal-rank fusion employed by the hybrid pipeline did not improve precision "
    "relative to single-method dense retrieval.")

R["The vanilla LLM baseline exhibits a hallucination rate of 0.58"] = (
    "The vanilla baseline records the highest hallucination rate, 0.320, when its answers are assessed "
    "against the reference answers, indicating that approximately one third of its claims are unsupported "
    "by or contradicted by the reference. All four retrieval-augmented pipelines record lower hallucination "
    "rates: standard RAG and query reformulation are lowest at 0.248, followed by hybrid retrieval (0.280) "
    "and multi-query expansion (0.283). The reduction relative to the baseline, from 0.320 to 0.248 for the "
    "strongest configurations, indicates that retrieval grounding reduces unsupported content, although the "
    "magnitude of the reduction is modest at this sample scale. Hypothesis 2, which predicted that hybrid "
    "retrieval would achieve the lowest hallucination rate, is not supported, since hybrid retrieval records "
    "a higher rate than both standard RAG and query reformulation. It should be noted that the baseline "
    "hallucination rate is assessed against the reference answer, whereas the retrieval pipelines are "
    "assessed against their retrieved context, so the comparison is indicative rather than exact.")

R["Groundedness measures the alignment between generated answers and retrieved evidence"] = (
    "Two further indicators of output quality are reported alongside the hallucination rate. Groundedness "
    "is operationalised here as the complement of the hallucination rate (one minus the rate), and "
    "correctness is operationalised as the ROUGE-L F-score overlap between the generated answer and the "
    "MedQuAD reference answer. Table 5.4 reports the correctness (ROUGE-L) and heuristic safety-orientation "
    "scores, and the groundedness values are reported alongside the hallucination rate in Table 5.3.")

R["Several patterns emerge. First, all pipelines perform best on symptom questions"] = (
    "Faithfulness was disaggregated by question type for the four most frequent categories in the sample "
    "(information, treatment, causes, and symptoms). The differences across question types were small and "
    "did not follow a consistent ordering across pipelines, and several categories contained few questions "
    "(for example, nine causes and nine symptom questions), so the per-type estimates carry wide "
    "uncertainty. The most consistent observation is that the retrieval pipelines, and multi-query expansion "
    "in particular, tended to achieve their highest faithfulness on information- and symptom-type questions, "
    "where consumer-health documents provide direct factual answers, whereas treatment questions, which "
    "often involve patient-specific considerations, were answered with lower faithfulness across all "
    "configurations. The earlier expectation of a uniform hybrid-retrieval advantage is not borne out by "
    "these data.")

R["Safety compliance measures whether generated answers adhere to clinical safety guidelines"] = (
    "Safety orientation was assessed using a heuristic lexical score that rewards expressions of uncertainty "
    "and referral to professionals and penalises unqualified definitive statements. The scores were "
    "uniformly modest and did not differ markedly across configurations (0.282 for the vanilla baseline and "
    "0.194 to 0.205 for the retrieval pipelines). The marginally higher score for the baseline is "
    "attributable to its more frequent use of hedging language, whereas the retrieval pipelines, which "
    "report specific retrieved facts, are penalised by the heuristic despite consistently appending a "
    "recommendation to consult a qualified professional. The heuristic is therefore a coarse proxy that "
    "captures surface markers of caution rather than substantive clinical safety; a robust safety assessment "
    "would require expert human review or a dedicated safety classifier, which is identified as a limitation "
    "and as a direction for future work.")

R["Research Question 4 asks where failures mainly occur"] = (
    "Research Question 4 concerns where failures mainly occur. Four failure categories were operationalised "
    "from the retrieved context and the evaluation scores: missing evidence (the retrieved context lacked "
    "the information needed to answer the question), noisy evidence (the retrieved context contained "
    "substantial irrelevant material), unsupported claims (the answer asserted content absent from the "
    "context), and unsafe tone (definitive advice offered without appropriate caveats). Table 5.6 reports, "
    "for each pipeline, the percentage of answered questions exhibiting each pattern; a single question may "
    "exhibit more than one pattern.")

R["The vanilla baseline shows a relatively even distribution"] = (
    "The patterns differ sharply between the baseline and the retrieval pipelines. For the vanilla baseline, "
    "missing evidence and unsupported claims are recorded for essentially every question: because the "
    "baseline retrieves no context, every claim is, by construction, unsupported by retrieved evidence, so "
    "these two categories are definitional rather than informative for it. For the four retrieval pipelines, "
    "the dominant pattern is noisy evidence (82 to 92 per cent of answered questions), reflecting that the "
    "retrieved set frequently contained some marginally relevant or off-topic passages alongside the "
    "relevant ones; missing evidence was comparatively rare (2 to 10 per cent), and multi-query expansion, "
    "which retrieves over several query variants, recorded both the lowest missing-evidence rate (2 per cent) "
    "and the highest noisy-evidence rate (92 per cent), consistent with a coverage-for-noise trade-off. "
    "Unsupported claims and unsafe tone were rarely the primary failure for the retrieval pipelines. The "
    "predominance of noisy evidence indicates that, for these pipelines, the principal opportunity for "
    "improvement lies in sharpening retrieval precision rather than in increasing coverage.")

R["The vanilla baseline is the fastest at 850 milliseconds"] = (
    "Because the per-question latencies were right-skewed by occasional API retries, median latency is "
    "reported as a robust summary. The fastest configurations are hybrid retrieval and standard RAG, with "
    "median end-to-end latencies of 2.8 seconds, followed by query reformulation (3.6 s) and the vanilla "
    "baseline (3.9 s); the baseline is not the fastest because its single generation call is unconstrained "
    "by retrieved context and tends to produce longer answers. Multi-query expansion is the slowest at a "
    "median of 6.5 seconds, since it issues an expansion call and retrieves over several query variants "
    "before generation. All configurations therefore operate within a few seconds per query under API "
    "access; the absolute values include network round-trip and should be read as comparative rather than "
    "as deployment figures.")

R["To determine whether the observed differences between pipelines are statistically significant"] = (
    "To determine whether the observed differences between pipelines are statistically significant at the "
    "pilot scale, a one-way analysis of variance was conducted on the RAGAS faithfulness scores, with five "
    "groups of 50 observations each (N = 250). The analysis tests the null hypothesis that all five "
    "pipelines share equal population means, and the F-statistic is defined in Equation 5.1. Because the "
    "evaluation was conducted at a pilot scale, the result should be regarded as indicative, and "
    "confirmatory testing is recommended once the full-scale evaluation has been completed.")

R["The analysis yields a highly significant F-statistic"] = (
    "The analysis yields an F-statistic of 2.08 with a p-value of 0.084 (F(4, 245) = 2.08), which does not "
    "reach the conventional 0.05 threshold; the null hypothesis of equal faithfulness means across the five "
    "pipelines therefore cannot be rejected at this sample size. The numerical ordering of the pipelines is "
    "consistent with the descriptive findings reported above, but the substantial within-pipeline variance "
    "(the pooled standard deviation exceeds 0.40) and the modest between-pipeline differences mean that a "
    "larger sample is required to establish statistical significance. Given this outcome, post-hoc pairwise "
    "testing and effect-size estimation are not reported, as they would not be warranted by a non-significant "
    "omnibus test at pilot scale; both are deferred to the planned full-scale evaluation. A weak negative "
    "association between faithfulness and the hallucination rate is observed across the five pipeline means "
    "(Pearson r = -0.47), in the expected direction, though it is not statistically significant given only "
    "five aggregate points.")

R["Research Question 1 asked how the four retrieval-based pipelines differ"] = (
    "Research Question 1 concerned how the four retrieval-based pipelines differ from the vanilla baseline in "
    "reducing hallucination. The results indicate that all four record higher faithfulness (0.485 to 0.614 "
    "against 0.381) and lower hallucination rates (0.248 to 0.283 against 0.320) than the baseline; the "
    "direction of the effect supports the premise that retrieval grounding reduces unsupported generation, "
    "although the analysis of variance shows that the differences are not statistically significant at the "
    "pilot scale (Section 5.8). Research Question 2 concerned the effect of retrieval depth; because the "
    "number of retrieved documents was held fixed at five as a control variable, depth was not varied in the "
    "main comparison, and the effect of broader coverage was instead examined through multi-query expansion, "
    "which retrieved over several query variants and recorded the highest faithfulness and recall. Research "
    "Question 3 concerned how the strategies compare across question types; the per-question-type analysis "
    "found small and inconsistent differences with wide uncertainty owing to small per-type samples, with "
    "the retrieval pipelines tending to perform best on information- and symptom-type questions. Research "
    "Question 4 concerned where failures occur; the failure analysis found that the retrieval pipelines are "
    "dominated by noisy-evidence cases rather than missing evidence, indicating that improving retrieval "
    "precision is the principal opportunity for further gains.")

R["The finding that hybrid retrieval outperforms single-method retrieval aligns with Woelk"] = (
    "The finding that retrieval grounding improves faithfulness and reduces hallucination relative to a "
    "non-retrieval baseline is consistent with Woelk (2025), who reported that retrieval reduced clinical "
    "decision errors, and with the broader evidence that grounding mitigates hallucination (Lewis et al., "
    "2020; Gao et al., 2023). The present study does not, however, reproduce the common expectation that "
    "hybrid retrieval is uniformly superior: on this sample multi-query expansion achieved the highest "
    "faithfulness and recall, while standard RAG achieved the highest precision and, with query "
    "reformulation, the lowest hallucination. This is consistent with the observation of Kumari et al. "
    "(2026) that retrieval design interacts with question characteristics, and with the finding of Strich "
    "et al. (2026) that no single configuration is uniformly optimal. The result underlines that the "
    "advantage of a given retrieval strategy is contingent on the corpus and question mix rather than fixed.")

R["The results have several practical implications. First, the vanilla LLM baseline is unsuitable"] = (
    "The results carry several practical implications, framed cautiously given the pilot scale. First, the "
    "non-retrieval baseline is the weakest configuration for evidence-grounded answering, recording the "
    "lowest faithfulness and the highest hallucination rate; although it produces fluent and relevant "
    "answers, these are less well supported, which is undesirable in a clinical-information setting. Second, "
    "standard RAG offers an attractive balance of grounding and simplicity, recording the highest context "
    "precision, the joint-lowest hallucination rate, and the second-fastest median latency, and is therefore "
    "a sound default where implementation simplicity matters. Third, multi-query expansion is preferable "
    "where faithfulness and recall are paramount and the additional latency (a median of 6.5 seconds) is "
    "acceptable. Fourth, hybrid retrieval and query reformulation did not, on this sample, justify their "
    "additional complexity, although the differences between configurations were not statistically "
    "significant and a larger evaluation may separate them more clearly.")

R["This chapter has presented the quantitative results from the pilot-scale comparative evaluation. The key findings are that all four"] = (
    "This chapter has presented the quantitative results from the pilot-scale comparative evaluation. The "
    "key findings are that all four retrieval-augmented pipelines recorded higher faithfulness and lower "
    "hallucination rates than the non-retrieval baseline; that multi-query expansion recorded the highest "
    "faithfulness (0.614) and context recall, while standard RAG recorded the highest context precision "
    "(0.725) and, with query reformulation, the lowest hallucination rate (0.248); that hybrid retrieval did "
    "not outperform the simpler strategies, so the prediction of a uniform hybrid advantage was not "
    "supported; that the between-pipeline differences in faithfulness were not statistically significant at "
    "the pilot scale (F(4, 245) = 2.08, p = 0.084); and that the retrieval pipelines were dominated by "
    "noisy-evidence rather than missing-evidence failures. The discussion related these findings to the "
    "literature and derived cautious practical implications. The next chapter summarises the study, "
    "acknowledges its limitations, and proposes directions for future research.")

R["The first key finding is that retrieval-augmented generation substantially reduces hallucination"] = (
    "The first key finding is that retrieval-augmented generation reduced hallucination and improved "
    "faithfulness relative to the non-retrieval baseline: the baseline recorded a hallucination rate of "
    "0.320 and a faithfulness of 0.381, whereas the retrieval pipelines recorded hallucination rates of "
    "0.248 to 0.283 and faithfulness of 0.485 to 0.614. The second key finding is that no single retrieval "
    "strategy dominated: multi-query expansion recorded the highest faithfulness and recall, standard RAG "
    "the highest precision and (with query reformulation) the lowest hallucination, and hybrid retrieval did "
    "not lead on any primary metric, so the prediction that hybrid retrieval would be uniformly superior was "
    "not supported. The third key finding is that the between-pipeline differences, though in the expected "
    "direction, were not statistically significant at the pilot scale (F(4, 245) = 2.08, p = 0.084), so the "
    "results are indicative and require confirmation at full scale. The fourth key finding is a trade-off "
    "between grounding and fluency: the baseline produced the most relevant and lexically complete answers "
    "because it always responds, whereas the retrieval pipelines were more conservative and occasionally "
    "declined when evidence was insufficient. The fifth key finding is that the retrieval pipelines' "
    "dominant failure mode was noisy rather than missing evidence, indicating that sharper retrieval "
    "precision, rather than greater coverage, is the principal opportunity for improvement.")

# Equation 5.2 (Cohen's d) and its caption are removed (not reported at pilot scale)
R["d = (M₁ − M₂) / S_pooled"] = ""
R["Equation 5.2: Cohen's d effect size"] = ""

for prefix, newtext in R.items():
    replace_prefix(paras, prefix, newtext)

# ── FIGURE CAPTIONS ──────────────────────────────────────────────────────────
CAPS = {
    "Figure 5.1:": "Figure 5.1: RAGAS faithfulness by pipeline (means with 95 per cent confidence intervals). Multi-query expansion records the highest faithfulness, and all four retrieval pipelines exceed the non-retrieval baseline.",
    "Figure 5.2:": "Figure 5.2: RAGAS context precision and recall by pipeline. Standard RAG records the highest precision and multi-query expansion the highest recall; these metrics are not applicable to the non-retrieval baseline.",
    "Figure 5.3:": "Figure 5.3: DeepEval hallucination rate by pipeline (lower is better). Standard RAG and query reformulation record the lowest rates.",
    "Figure 5.4:": "Figure 5.4: RAGAS answer relevance by pipeline. The non-retrieval baseline records the highest relevance because it always produces a complete answer, whereas the retrieval pipelines occasionally decline when evidence is insufficient.",
    "Figure 5.5:": "Figure 5.5: Faithfulness by question type and pipeline for the four most frequent categories. Per-type samples are small, so the estimates carry wide uncertainty.",
    "Figure 5.6:": "Figure 5.6: Heuristic safety-orientation score by pipeline. The scores are uniformly modest and are not strongly discriminative (see Section 5.5).",
    "Figure 5.7:": "Figure 5.7: Primary failure-pattern distribution by pipeline. The retrieval pipelines are dominated by noisy-evidence cases, whereas the baseline, lacking retrieved context, is recorded as missing evidence by construction.",
    "Figure 5.8:": "Figure 5.8: Median end-to-end latency by pipeline. Multi-query expansion is the slowest, owing to multiple retrieval passes; the remaining pipelines are comparable.",
    "Figure 5.9:": "Figure 5.9: Relationship between RAGAS faithfulness and the DeepEval hallucination rate across the five pipeline means (Pearson r = -0.47; not significant at this scale).",
}
for pref, txt in CAPS.items():
    replace_prefix(paras, pref, txt)

# ── TABLE CAPTIONS that need correcting ──────────────────────────────────────
replace_prefix(paras, "Table 5.7:", "Table 5.7: Median end-to-end latency by pipeline (seconds) and latency relative to the fastest configuration.")
replace_prefix(paras, "Table 5.4:", "Table 5.4: Correctness (ROUGE-L F-score) and heuristic safety-orientation scores by pipeline, with standard deviations.")
replace_prefix(paras, "Table 4.2:", "Table 4.2: Configuration settings selected for the main experiment, with rationale.")

d.save(PATH)
print("Prose, captions and global config fixes applied.")
