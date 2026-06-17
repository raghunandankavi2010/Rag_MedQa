#!/usr/bin/env python3
"""Apply SAFE, results-independent appendix edits to the thesis .docx.

Scope (no API / no experimental data required):
  - Populate Appendix E (Glossary), currently empty.
  - Correct Appendix F output-file references to the real repository filenames
    and insert the real video-presentation link.
  - Replace the placeholder notes in Appendix A (Research Proposal) and
    Appendix B (Ethics Approval) with proper referencing text.

Results-coupled edits (Chapter 5 numbers, methodology scale, ANOVA, Appendix D
tables, figures) are intentionally NOT touched here; they are applied in a
separate pass once the real comparison run has produced data.
"""
import docx

PATH = "/Users/raghunandan.k/Downloads/RaghunandanKavi_PN1196933_RAG_Medical_QA_FINAL.docx"

GLOSSARY = [
    ("Retrieval-Augmented Generation (RAG)", "An architecture in which language-model generation is grounded in documents retrieved from an external corpus at inference time, reducing reliance on parametric memory."),
    ("Large Language Model (LLM)", "A transformer-based model trained on large text corpora to understand and generate natural language."),
    ("Hallucination", "The generation of content that is unsupported by, or contradictory to, the source evidence or established fact."),
    ("Faithfulness (RAGAS)", "The proportion of claims in a generated answer that are supported by the retrieved context."),
    ("Answer Relevance (RAGAS)", "The degree to which a generated answer addresses the intent of the original question."),
    ("Context Precision (RAGAS)", "The proportion of retrieved context that is relevant to the question."),
    ("Context Recall (RAGAS)", "The extent to which the retrieved context contains the information needed to answer the question."),
    ("Groundedness (DeepEval)", "The alignment between a generated answer and the retrieved evidence."),
    ("Hallucination Rate (DeepEval)", "The proportion of unsupported or fabricated claims in a generated answer."),
    ("Safety Compliance", "Adherence to clinical safety guidelines, including expression of uncertainty when evidence is insufficient and encouragement of professional consultation."),
    ("Dense Retrieval", "Retrieval based on similarity between learned vector embeddings of the query and the documents."),
    ("Sparse Retrieval", "Lexical retrieval based on term overlap, such as BM25 or TF-IDF."),
    ("BM25", "A probabilistic ranking function that scores documents by weighted term frequency for sparse keyword retrieval."),
    ("TF-IDF", "Term Frequency-Inverse Document Frequency; a weighting scheme reflecting term importance in a document relative to a corpus."),
    ("Reciprocal Rank Fusion (RRF)", "A rank-based method that fuses multiple ranked lists by summing reciprocal ranks, robust to differences in score scale."),
    ("Hybrid Retrieval", "Retrieval that combines dense and sparse methods, typically fused with Reciprocal Rank Fusion."),
    ("Multi-Query Expansion", "Generation of several paraphrased queries to broaden retrieval coverage before merging the retrieved results."),
    ("Query Reformulation", "Rewriting a user query into a clearer, retrieval-friendly form prior to search."),
    ("Cross-Encoder Re-ranker", "A model that jointly encodes a query and a candidate passage to produce a sharp relevance score for re-ordering candidates."),
    ("Chunking", "The process of splitting documents into fixed-size segments, with overlap, for embedding and retrieval."),
    ("Embedding", "A dense vector representation of text within a semantic space."),
    ("Vector Database (ChromaDB)", "A store for embeddings supporting approximate nearest-neighbour similarity search; ChromaDB is the implementation used in this study."),
    ("Top-k", "The number of highest-ranked documents retrieved for a query."),
    ("Temperature", "A decoding parameter controlling generation randomness; set to 0 for deterministic output."),
    ("Safety Guardrail", "A pre-retrieval classifier that blocks unsafe or off-domain queries before answer generation."),
    ("Context Compression", "Summarisation of retrieved passages to the most relevant content before generation, mitigating long-context degradation."),
    ("MedQuAD", "The Medical Question Answering Dataset, comprising 47,457 question-answer pairs drawn from twelve National Institutes of Health sources."),
    ("RAGAS", "Retrieval-Augmented Generation Assessment; an LLM-based evaluation framework measuring faithfulness, answer relevance, and context precision/recall."),
    ("DeepEval", "An evaluation framework assessing language-model output reliability, hallucination, and safety behaviour."),
    ("ANOVA", "Analysis of Variance; a statistical test for differences among the means of three or more groups."),
]

APPENDIX_F_FILES = [
    "scripts/compare_pipelines.py - constructs and evaluates the five pipeline configurations under controlled conditions.",
    "scripts/eda.py - exploratory data analysis of the MedQuAD corpus (summary statistics and plots).",
    "notebooks/rag_corrected.py - the integrated retrieval pipeline (guardrail, hybrid retrieval, re-ranking, compression).",
    "reports/pipeline_comparison_summary.csv - aggregate RAGAS, DeepEval and ROUGE scores per pipeline.",
    "reports/pipeline_comparison_raw.csv - per-question, per-pipeline metrics and generated answers.",
    "reports/pipeline_comparison_anova.txt - one-way ANOVA on faithfulness across the pipelines.",
    "reports/elite_rag_evaluation.csv - per-question evaluation of the integrated pipeline.",
    "reports/eda_summary.csv and reports/eda_qtype_distribution.csv - dataset EDA outputs.",
]
VIDEO_LINK = "https://drive.google.com/file/d/1pAxQn5kUAXREogXSsIaCqE-YUeStF1of/view?usp=drive_link"


def set_text(p, text):
    """Replace a paragraph's text, preserving the first run's formatting."""
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


def find_para(ps, predicate):
    for p in ps:
        if predicate(p):
            return p
    return None


def main():
    d = docx.Document(PATH)
    ps = d.paragraphs

    # ── Appendix E: Glossary ────────────────────────────────────────────────
    appF = find_para(ps, lambda p: p.text.strip().startswith("Appendix F:"))
    appE = find_para(ps, lambda p: p.text.strip().startswith("Appendix E:"))
    # Intro line: the (empty) paragraph immediately after the Appendix E heading
    idxE = ps.index(appE)
    intro = ps[idxE + 1]
    set_text(intro, "The following terms and abbreviations are used throughout this thesis.")
    for term, definition in GLOSSARY:
        gp = appF.insert_paragraph_before(style="Normal")
        rt = gp.add_run(f"{term} - ")
        rt.bold = True
        gp.add_run(definition)
    print(f"Inserted {len(GLOSSARY)} glossary entries before Appendix F.")

    # ── Appendix F: correct output-file references + video link ─────────────
    ps = d.paragraphs  # refresh after insertion
    # Replace the three old bullet references (thesis_experiment_results.csv etc.)
    old_bullets = [p for p in ps if p.style.name == "List Bullet"
                   and "thesis_" in p.text and ".csv" in p.text or
                   (p.style.name == "List Bullet" and "thesis_final_report" in p.text)]
    # More robust: gather the contiguous bullet block mentioning thesis_ outputs
    bullets = [p for p in ps if p.style.name == "List Bullet" and "thesis_" in p.text]
    for i, p in enumerate(bullets):
        if i < len(APPENDIX_F_FILES):
            set_text(p, APPENDIX_F_FILES[i])
    # Append any remaining real files as new bullets after the last one
    if bullets:
        anchor = bullets[-1]
        for extra in APPENDIX_F_FILES[len(bullets):]:
            np = anchor.insert_paragraph_before(style="List Bullet")
            set_text(np, extra)
    # Video link line
    vid = find_para(d.paragraphs, lambda p: "Video presentation" in p.text or "Thesis defence presentation" in p.text)
    if vid:
        set_text(vid, f"Thesis defence presentation (video): {VIDEO_LINK}")
    # Soften the trailing note about generated outputs
    note = find_para(d.paragraphs, lambda p: p.text.strip().startswith("Note: the output CSV files are generated"))
    if note:
        set_text(note, "Note: the comparison output files (pipeline_comparison_*.csv and the ANOVA summary) "
                       "are produced when scripts/compare_pipelines.py is executed against the held-out "
                       "evaluation set; the integrated-pipeline and EDA outputs are already included in the repository.")
    print("Updated Appendix F file references, video link, and note.")

    # ── Appendix A: Research Proposal ───────────────────────────────────────
    a_link = find_para(d.paragraphs, lambda p: p.text.strip().startswith("Research proposal:"))
    if a_link:
        set_text(a_link, "The approved research proposal is provided as a separately submitted document "
                         "(RaghunandanKavi_PN1196933_OriginalProposal), which defined the aim, objectives, "
                         "research questions, and indicative methodology subsequently refined in this thesis.")
    a_note = find_para(d.paragraphs, lambda p: p.text.strip().startswith("Note: replace the placeholder link above with the shared location of the approved proposal"))
    if a_note:
        set_text(a_note, "")
    print("Updated Appendix A.")

    # ── Appendix B: Ethics Approval ─────────────────────────────────────────
    b_link = find_para(d.paragraphs, lambda p: p.text.strip().startswith("Ethics approval:"))
    if b_link:
        set_text(b_link, "The thesis supervisor's written approval of the research proposal is provided as a "
                         "separately submitted document (Thesis Supervisor Approval). As the study used only the "
                         "publicly available MedQuAD dataset and commercial language-model APIs, with no human "
                         "participants or personally identifiable information, no further ethical clearance was required.")
    b_note = find_para(d.paragraphs, lambda p: p.text.strip().startswith("Note: replace the placeholder link above with the shared location of the signed approval"))
    if b_note:
        set_text(b_note, "")
    print("Updated Appendix B.")

    d.save(PATH)
    print(f"\nSaved -> {PATH}")


if __name__ == "__main__":
    main()
