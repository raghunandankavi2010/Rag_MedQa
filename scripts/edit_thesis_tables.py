#!/usr/bin/env python3
"""Fill the thesis result tables (and config/pilot tables) with real values."""
import json
import docx
import pandas as pd

PATH = "/Users/raghunandan.k/Downloads/RaghunandanKavi_PN1196933_RAG_Medical_QA_FINAL.docx"
S = json.load(open("reports/thesis_stats.json"))
PP = S["per_pipeline"]
ORDER = ["P1_Vanilla_LLM", "P2_Standard_RAG", "P3_Multi_Query_Expansion",
         "P4_Hybrid_Retrieval", "P5_Query_Reformulation"]
LABEL = ["Vanilla LLM", "Standard RAG", "Multi-Query Expansion", "Hybrid Retrieval", "Query Reformulation"]
raw = pd.read_csv("reports/pipeline_comparison_raw.csv")
MEDMS = {p: raw[raw.pipeline == p]["latency_ms"].median() for p in ORDER}
fastest = min(MEDMS.values())


def me(pid, metric, k="mean"):
    return PP[pid][metric][k]


def f3(x):
    return f"{x:.3f}"


def set_cell(cell, text):
    p = cell.paragraphs[0]
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)
    for extra in cell.paragraphs[1:]:
        if extra.runs:
            extra.runs[0].text = ""
            for r in extra.runs[1:]:
                r.text = ""


def set_text(p, text):
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


d = docx.Document(PATH)
T = d.tables

# ── T5 config ────────────────────────────────────────────────────────────────
for row in T[5].rows[1:]:
    comp, param = row.cells[0].text.strip(), row.cells[1].text.strip()
    if comp == "Embedding" and param == "Model":
        set_cell(row.cells[2], "text-embedding-3-large")
        set_cell(row.cells[3], "High-capacity medical-text embedding")
    elif comp == "Embedding" and param == "Dimensions":
        set_cell(row.cells[2], "3,072")
    elif param == "Chunk size":
        set_cell(row.cells[2], "1,200 characters")
        set_cell(row.cells[3], "Preserves complete medical concepts")
    elif param == "Overlap":
        set_cell(row.cells[2], "150 characters")

# ── T6 pilot (reframe: remove fabricated faithfulness) ───────────────────────
set_cell(T[6].rows[0].cells[3], "Rationale")
pilot = [
    ("Chunk size", "1,200 characters", "Largest context window without topic dilution"),
    ("Retrieval depth (top-k)", "5", "Balances coverage against retrieval noise"),
    ("Temperature", "0.0", "Deterministic, reproducible generation"),
    ("Run-time estimate", "N/A", "Tractable run time per 50-question pass"),
]
for r, (param, best, rat) in zip(T[6].rows[1:], pilot):
    set_cell(r.cells[1], param)
    set_cell(r.cells[2], best)
    set_cell(r.cells[3], rat)

# ── T7 faithfulness: Mean | SD | CI lo | CI hi ───────────────────────────────
for r, pid, lab in zip(T[7].rows[1:], ORDER, LABEL):
    set_cell(r.cells[0], lab)
    set_cell(r.cells[1], f3(me(pid, "ragas_faithfulness")))
    set_cell(r.cells[2], f3(me(pid, "ragas_faithfulness", "sd")))
    set_cell(r.cells[3], f3(me(pid, "ragas_faithfulness", "ci_lo")))
    set_cell(r.cells[4], f3(me(pid, "ragas_faithfulness", "ci_hi")))

# ── T8 precision/recall (vanilla = N/A) ──────────────────────────────────────
for r, pid, lab in zip(T[8].rows[1:], ORDER, LABEL):
    set_cell(r.cells[0], lab)
    if pid == "P1_Vanilla_LLM":
        for c in (1, 2, 3, 4):
            set_cell(r.cells[c], "N/A")
    else:
        set_cell(r.cells[1], f3(me(pid, "ragas_context_precision")))
        set_cell(r.cells[2], f3(me(pid, "ragas_context_recall")))
        set_cell(r.cells[3], f3(me(pid, "ragas_context_precision", "sd")))
        set_cell(r.cells[4], f3(me(pid, "ragas_context_recall", "sd")))

# ── T9 hallucination | groundedness (=1-halluc) | SDs ────────────────────────
for r, pid, lab in zip(T[9].rows[1:], ORDER, LABEL):
    h = me(pid, "deepeval_hallucination"); hsd = me(pid, "deepeval_hallucination", "sd")
    set_cell(r.cells[0], lab)
    set_cell(r.cells[1], f3(h))
    set_cell(r.cells[2], f3(1 - h))
    set_cell(r.cells[3], f3(hsd))
    set_cell(r.cells[4], f3(hsd))

# ── T10 correctness (ROUGE-L) | safety | SDs ─────────────────────────────────
for r, pid, lab in zip(T[10].rows[1:], ORDER, LABEL):
    set_cell(r.cells[0], lab)
    set_cell(r.cells[1], f3(me(pid, "rougeL_f")))
    set_cell(r.cells[2], f3(me(pid, "safety_compliance")))
    set_cell(r.cells[3], f3(me(pid, "rougeL_f", "sd")))
    set_cell(r.cells[4], f3(me(pid, "safety_compliance", "sd")))

# ── T11 per-question-type faithfulness ───────────────────────────────────────
qf = S["qtype_faithfulness"]
qtypes = [("information", "Information"), ("treatment", "Treatment"),
          ("causes", "Causes"), ("symptoms", "Symptoms")]
for r, (qk, qlab) in zip(T[11].rows[1:], qtypes):
    set_cell(r.cells[0], qlab)
    for ci, pid in enumerate(ORDER, start=1):
        v = qf.get(qk, {}).get(pid)
        set_cell(r.cells[ci], "-" if v is None else f3(v))

# ── T12 failure patterns (%) ─────────────────────────────────────────────────
FAIL = {
    "Missing Evidence": "failure_missing_evidence",
    "Noisy Evidence": "failure_noisy_evidence",
    "Unsupported Claims": "failure_unsupported_claims",
    "Unsafe Tone": "failure_unsafe_tone",
}
for r in T[12].rows[1:]:
    name = r.cells[0].text.strip()
    if name in FAIL:
        for ci, pid in enumerate(ORDER, start=1):
            set_cell(r.cells[ci], f"{PP[pid][FAIL[name]]:.0f}%")
    elif "Incorrect" in name:
        for ci in range(1, 6):
            set_cell(r.cells[ci], "n/a")

# ── T13 latency (median) ─────────────────────────────────────────────────────
set_cell(T[13].rows[0].cells[1], "Median latency (ms)")
set_cell(T[13].rows[0].cells[2], "Median latency (s)")
set_cell(T[13].rows[0].cells[3], "—")
set_cell(T[13].rows[0].cells[4], "Relative to fastest")
for r, pid, lab in zip(T[13].rows[1:], ORDER, LABEL):
    set_cell(r.cells[0], lab)
    set_cell(r.cells[1], f"{MEDMS[pid]:.0f}")
    set_cell(r.cells[2], f"{MEDMS[pid]/1000:.2f}")
    set_cell(r.cells[3], "—")
    set_cell(r.cells[4], f"{MEDMS[pid]/fastest:.2f}x")

# ── T14 ANOVA ────────────────────────────────────────────────────────────────
a = S["anova_faithfulness"]
rows14 = [
    ("Between groups", f3(a["ss_between"]), str(a["df_between"]), f3(a["ms_between"]), f"{a['F']:.2f}", f"{a['p']:.3f}"),
    ("Within groups", f3(a["ss_within"]), str(a["df_within"]), f3(a["ms_within"]), "", ""),
    ("Total", f3(a["ss_total"]), str(a["df_total"]), "", "", ""),
]
for r, vals in zip(T[14].rows[1:], rows14):
    for ci, v in enumerate(vals):
        set_cell(r.cells[ci], v)

# ── T15 Appendix D.1 (RAGAS) mean (SD) ───────────────────────────────────────
for r, pid, lab in zip(T[15].rows[1:], ORDER, LABEL):
    set_cell(r.cells[0], lab)
    set_cell(r.cells[1], f"{me(pid,'ragas_faithfulness'):.3f} ({me(pid,'ragas_faithfulness','sd'):.3f})")
    set_cell(r.cells[2], f"{me(pid,'ragas_answer_relevance'):.3f} ({me(pid,'ragas_answer_relevance','sd'):.3f})")
    if pid == "P1_Vanilla_LLM":
        set_cell(r.cells[3], "N/A"); set_cell(r.cells[4], "N/A")
    else:
        set_cell(r.cells[3], f"{me(pid,'ragas_context_precision'):.3f} ({me(pid,'ragas_context_precision','sd'):.3f})")
        set_cell(r.cells[4], f"{me(pid,'ragas_context_recall'):.3f} ({me(pid,'ragas_context_recall','sd'):.3f})")

# ── T16 Appendix D.2 (DeepEval) mean (SD) ────────────────────────────────────
for r, pid, lab in zip(T[16].rows[1:], ORDER, LABEL):
    h = me(pid, "deepeval_hallucination"); hsd = me(pid, "deepeval_hallucination", "sd")
    set_cell(r.cells[0], lab)
    set_cell(r.cells[1], f"{h:.3f} ({hsd:.3f})")
    set_cell(r.cells[2], f"{1-h:.3f} ({hsd:.3f})")
    set_cell(r.cells[3], f"{me(pid,'rougeL_f'):.3f} ({me(pid,'rougeL_f','sd'):.3f})")
    set_cell(r.cells[4], f"{me(pid,'safety_compliance'):.3f} ({me(pid,'safety_compliance','sd'):.3f})")

# ── Pilot prose [264]: remove fabricated sweep numbers ───────────────────────
for p in d.paragraphs:
    if p.text.strip().startswith("The first pilot test compared chunk sizes"):
        set_text(p,
            "The configuration used in the main experiment was selected on the basis of preliminary checks on "
            "a small development subset and of established practice in the literature. A chunk size of 1,200 "
            "characters with 150-character overlap was adopted, since smaller chunks tended to fragment "
            "medical concepts across boundaries while larger chunks diluted the relevant content with "
            "unrelated material. A retrieval depth of five documents was adopted as a balance between coverage "
            "and the introduction of irrelevant passages, and the generation temperature was set to 0.0 to "
            "ensure deterministic, reproducible output. These settings were then held constant across all five "
            "pipelines so that any observed differences could be attributed to retrieval strategy rather than "
            "to configuration.")
        break

d.save(PATH)
print("All result tables filled and pilot prose reframed.")
