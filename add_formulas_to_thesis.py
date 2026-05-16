"""
Add mathematical formulas to thesis document.
Inserts equations for cosine similarity, TF-IDF, RRF, and evaluation metrics.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document(r'C:\Users\HP\Downloads\RAG_Medical_QA_Thesis_LJMU_COMPLETE.docx')

def add_formula_paragraph(para, formula_text, caption_text):
    """Add a centered formula with caption after the given paragraph."""
    # Insert formula paragraph after current
    new_p = para._element
    
    formula_p = doc.add_paragraph()
    formula_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = formula_p.add_run(formula_text)
    run.font.size = Pt(12)
    run.font.name = 'Cambria Math'
    run.font.italic = True
    
    new_p.addnext(formula_p._element)
    
    if caption_text:
        cap_p = doc.add_paragraph()
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap_p.add_run(caption_text)
        run.font.size = Pt(10)
        run.font.italic = True
        run.font.color.rgb = RGBColor(80, 80, 80)
        formula_p._element.addnext(cap_p._element)
    
    return formula_p

# Find key paragraphs to insert formulas after
formula_insertions = [
    # Cosine Similarity - after dense retrieval description
    {
        'search': "Dense retrieval captures semantic similarity between queries and documents, enabling matches on paraphrased or conceptually related terms.",
        'formula': "cos(θ) = (A · B) / (||A|| × ||B||) = (Σᵢ AᵢBᵢ) / (√(Σᵢ Aᵢ²) × √(Σᵢ Bᵢ²))",
        'caption': "Equation 3.1: Cosine similarity between query embedding A and document embedding B",
    },
    # TF-IDF - after sparse retrieval description
    {
        'search': "Sparse retrieval uses keyword matching to ensure that exact medical terminology is preserved.",
        'formula': "TF-IDF(t, d, D) = tf(t, d) × idf(t, D) = fₜ,ₐ / Σₖ fₖ,ₐ × log(|D| / |{d ∈ D : t ∈ d}|)",
        'caption': "Equation 3.2: TF-IDF weighting for term t in document d across corpus D",
    },
    # RRF - after hybrid retrieval description
    {
        'search': "Reciprocal rank fusion was selected over linear interpolation because it handles rank information directly rather than requiring score normalisation.",
        'formula': "RRF(d) = Σᵣ ₌ ₁ⁿ 1 / (k + rankᵣ(d))    where k = 60",
        'caption': "Equation 3.3: Reciprocal Rank Fusion score for document d across n retrieval methods",
    },
    # Faithfulness - after RAGAS description
    {
        'search': "Faithfulness measures whether the claims in the generated answer can be inferred from the retrieved context.",
        'formula': "F = (1/N) Σᵢ ₌ ₁ᴺ verify(claimᵢ, context)    where verify ∈ {0, 1}",
        'caption': "Equation 3.4: Faithfulness score as proportion of supported claims",
    },
    # Hallucination Rate - after DeepEval description
    {
        'search': "Hallucination rate measures the proportion of unsupported or fabricated claims in the generated answer.",
        'formula': "H = (1/N) Σᵢ ₌ ₁ᴺ (1 - verify(claimᵢ, context)) = 1 - F",
        'caption': "Equation 3.5: Hallucination rate as complement of faithfulness",
    },
    # Safety Score - after safety compliance description
    {
        'search': "Safety compliance evaluates whether the generated answer adheres to safety guidelines.",
        'formula': "S = Σₚ ₐ ₛ ₐ ₗ ₘ ₐ ₓ(0, 0.1 × I(present)) - Σₙ ₐ ₓ(0, 0.15 × I(present))",
        'caption': "Equation 3.6: Safety compliance score based on presence of uncertainty phrases and absence of definitive patterns",
    },
    # Word Overlap - after evaluation description
    {
        'search': "Word overlap measures the proportion of words in the ground truth that also appear in the generated answer.",
        'formula': "WO = |W_GT ∩ W_ANS| / |W_GT|",
        'caption': "Equation 3.7: Word overlap between ground truth and generated answer",
    },
    # ANOVA F-statistic - after statistical testing description
    {
        'search': "To determine whether observed differences between pipelines are statistically significant, one-way analysis of variance is conducted on RAGAS faithfulness scores.",
        'formula': "F = MSB / MSW = (SSB / (k - 1)) / (SSW / (N - k))",
        'caption': "Equation 5.1: One-way ANOVA F-statistic where MSB is mean square between groups, MSW is mean square within groups",
    },
    # Cohen's d - after effect size description
    {
        'search': "The largest effect size is between the vanilla baseline and hybrid retrieval",
        'formula': "d = (M₁ - M₂) / SD_pooled    where SD_pooled = √(((n₁ - 1)SD₁² + (n₂ - 1)SD₂²) / (n₁ + n₂ - 2))",
        'caption': "Equation 5.2: Cohen's d effect size for pairwise pipeline comparisons",
    },
]

inserted_count = 0
for item in formula_insertions:
    found = False
    for para in doc.paragraphs:
        if item['search'] in para.text:
            add_formula_paragraph(para, item['formula'], item['caption'])
            inserted_count += 1
            found = True
            break
    if not found:
        print(f"WARNING: Could not find text for: {item['search'][:60]}...")

# Save
output_path = r'C:\Users\HP\Downloads\RAG_Medical_QA_Thesis_LJMU_FINAL_WITH_FORMULAS.docx'
doc.save(output_path)
print(f"Added {inserted_count} formulas to thesis.")
print(f"Saved to: {output_path}")
